import os
import time

from win32com import client as wc
from retry import retry
from docx import Document


class WordDealer(object):
    def __init__(self, doc_path, key_word):
        self.doc_path = doc_path
        self.is_doc = False
        if '第' in key_word:
            self.key_word = key_word.split('第')[0]
        else:
            self.key_word = key_word
        if doc_path.endswith('.doc') and not doc_path.startswith('~$'):
            self.is_doc = True
            print('当前文件为doc格式，正在转换为docx格式用于处理', doc_path)

            self.doc2docx()

    def extract_paragraph(self):
        document = Document(self.doc_path)
        if not document.paragraphs:
            return []

        para = ''

        try:
            province = document.paragraphs[0].text
        except:
            province = ''

        try:
            completion_year = document.paragraphs[2].text
        except:
            completion_year = ''

        for paragraph in document.paragraphs:
            if self.key_word in paragraph.text:
                para += '\n' + paragraph.text

        return completion_year, province, para

    @retry(tries=10)
    def doc2docx(self):
        word = wc.Dispatch('Word.Application')

        doc = word.Documents.Open(os.path.abspath(self.doc_path))  # 目标路径下的文件
        self.doc_path += 'x'
        doc.SaveAs((os.path.abspath(self.doc_path)), 12, False, "", True, "", False, False, False,
                   False)  # 转化后路径下的文件
        doc.Close()
        word.Quit()
        print('.doc转换完成', self.doc_path)
        time.sleep(0.5)
